"""
DOCX 文档解析器 — 单一职责：将 DOCX 解析为增强文本（文本 + HTML 表格 + 图片描述/URL）。

使用 python-docx 提取段落文本、表格和内嵌图片：
    - 文本：doc.paragraphs（按文档顺序）
    - 表格：doc.tables → HTML <table>
    - 图片：inline_shapes / related_parts → 上传 MinIO / VLM 描述

分页检测（DOCX_PAGE_BREAK_DETECTION=True）：
    - 检测 <w:br w:type="page"/> — 用户插入的显式分页符；
    - 检测 <w:lastRenderedPageBreak/> — Word 渲染的分页位置提示；
    - 检测到分页符时递增页码，使 sections 携带真实页码信息。

通过遍历 doc.element.body 保持原始文档顺序（段落与表格交错出现），
而非先提取所有段落后提取所有表格。

遵循优雅降级：
    - python-docx 未安装 → 返回空字符串，调用方降级；
    - VLM 不可用 → 跳过图片描述，只保留文本和表格；
    - MinIO 不可用 → 跳过图片上传，降级为 VLM 描述；
    - 配置开关可单独关闭表格提取或图片提取。
"""

from __future__ import annotations

import asyncio
from typing import Any

from app.config import get_settings
from app.document.base import DocumentParser, ParsedSection
from app.utils.logger import get_logger

log = get_logger(__name__)

# VLM 并发控制
_VLM_SEMAPHORE_LIMIT: int = 3
_IMAGE_PROMPT: str = (
    "请用一句话描述这张图片的内容，"
    "重点关注图表、数据、文字和关键信息，便于后续检索。"
)

# DOCX XML 命名空间
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


class DOCXParser(DocumentParser):
    """DOCX 解析器 — python-docx 文本 + 表格 + 图片 URL/VLM + 分页检测。"""

    async def parse(self, file_path: str) -> str:
        """解析 DOCX 文档，返回增强文本。

        Args:
            file_path: DOCX 文件路径。

        Returns:
            增强文本（纯文本 + HTML 表格 + 图片描述/URL）。
            python-docx 未安装或解析失败时返回空字符串。
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            log.warning("docx.parser_skipped", reason="python-docx not installed")
            return ""

        try:
            docx_doc = DocxDocument(file_path)
        except Exception as exc:
            log.warning("docx.open_failed", file_path=file_path, error=str(exc))
            return ""

        settings = get_settings()
        table_enabled = self._bool(getattr(settings, "DOCX_TABLE_EXTRACTION_ENABLED", True), True)
        image_enabled = self._bool(getattr(settings, "DOCX_IMAGE_EXTRACTION_ENABLED", True), True)
        max_images = self._int(getattr(settings, "DOCX_IMAGE_MAX_PER_DOC", 50), 50)
        image_upload_enabled = self._bool(getattr(settings, "DOCX_IMAGE_UPLOAD_ENABLED", False), False)
        image_min_size = self._int(getattr(settings, "DOCX_IMAGE_MIN_SIZE", 50), 50)
        page_break_detection = self._bool(
            getattr(settings, "DOCX_PAGE_BREAK_DETECTION", True), True
        )

        sections: list[ParsedSection] = []
        image_count = 0

        # 遍历文档 body 元素，保持段落与表格的原始顺序
        current_page = 0  # 当前页码（从 0 开始）
        for element in docx_doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                # 段落 — 检测分页符
                if page_break_detection:
                    has_break = self._has_page_break(element)
                    if has_break and current_page > 0:
                        current_page += 1

                # 格式化段落 — 标题/列表/正文分别输出 HTML 标签
                content = self._format_paragraph(element)
                if content and content.strip():
                    sections.append(
                        ParsedSection(
                            kind="text",
                            content=content.strip(),
                            page=current_page,
                        )
                    )
            elif tag == "tbl" and table_enabled:
                # 表格
                html = self._extract_table_html(element)
                if html:
                    sections.append(
                        ParsedSection(
                            kind="table",
                            content=html,
                            page=current_page,
                        )
                    )

        # 提取图片（独立于 body 遍历，因为图片嵌在段落中）
        if image_enabled and image_count < max_images:
            img_sections, img_count = await self._extract_images(
                docx_doc,
                max_images - image_count,
                image_upload_enabled=image_upload_enabled,
                image_min_size=image_min_size,
            )
            sections.extend(img_sections)
            image_count += img_count

        # 提取页眉页脚
        header_footer_sections = self._extract_headers_footers(docx_doc)
        sections.extend(header_footer_sections)

        log.info(
            "docx.parsed",
            file_path=file_path,
            sections=len(sections),
            tables=sum(1 for s in sections if s.kind == "table"),
            images=sum(
                1 for s in sections if s.kind in ("image_desc", "image_url")
            ),
            pages=current_page + 1,
        )

        # 使用配置中的分页分隔符
        page_sep_raw = getattr(settings, "PAGE_SEPARATOR", "")
        page_separator = page_sep_raw if isinstance(page_sep_raw, str) else ""
        return self.sections_to_text(
            sections,
            page_separator=page_separator,
        )

    @staticmethod
    def _has_page_break(element: Any) -> bool:
        """检测段落是否包含分页符。

        检测两种分页标记：
            - <w:br w:type="page"/> — 用户显式插入的分页符；
            - <w:lastRenderedPageBreak/> — Word 上次渲染的分页位置。

        Args:
            element: <w:p> XML 元素。

        Returns:
            True 如果段落包含分页符。
        """
        try:
            # 方法 1: 遍历子元素查找 <w:br type="page">
            for child in element.iter():
                child_tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                if child_tag == "br":
                    br_type = child.get(f"{{{_W_NS}}}type", "")
                    if br_type == "page":
                        return True
                # 方法 2: 查找 <w:lastRenderedPageBreak/>
                if child_tag == "lastRenderedPageBreak":
                    return True
            return False
        except Exception:
            return False

    @staticmethod
    def _format_paragraph(element: Any) -> str:
        """从 <w:p> XML 元素格式化段落为 HTML — 保留标题层级和列表结构。

        样式映射规则（对齐 mammoth 的语义化映射）：
            - Title / 标题 / Heading 1 / 标题 1 → <h1>
            - Heading 2 / 标题 2 → <h2>
            - Heading 3 / 标题 3 → <h3>
            - Heading 4 / 标题 4 → <h4>
            - 列表（含 numPr）→ <ul><li>...</li></ul>
            - 其他（正文）→ 纯文本

        Args:
            element: <w:p> XML 元素。

        Returns:
            HTML 字符串。空字符串表示空段落。
        """
        try:
            # 直接从 XML 提取文本，避免 Paragraph(element, None) 的 parent 依赖问题
            text = DOCXParser._extract_text_from_element(element)
            text = text.strip()
            if not text:
                return ""

            # 1. 检测标题样式
            heading_html = DOCXParser._style_to_heading(element, text)
            if heading_html:
                return heading_html

            # 2. 检测列表（numPr）
            if DOCXParser._is_list_paragraph(element):
                return f"<ul><li>{DOCXParser._escape_html(text)}</li></ul>"

            # 3. 普通正文段落
            return DOCXParser._escape_html(text)

        except Exception:
            return ""

    @staticmethod
    def _extract_text_from_element(element: Any) -> str:
        """从 <w:p> XML 元素直接提取所有 <w:t> 文本内容。

        不依赖 python-docx 的 Paragraph 类（避免 parent=None 时 .text 为空的问题），
        直接遍历 XML 树提取 <w:t> 节点的文本。

        Args:
            element: <w:p> XML 元素。

        Returns:
            拼接后的纯文本。
        """
        try:
            ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
            texts: list[str] = []
            for t_node in element.iter(f"{ns}t"):
                if t_node.text:
                    texts.append(t_node.text)
            return "".join(texts)
        except Exception:
            return ""

    @staticmethod
    def _style_to_heading(element: Any, text: str) -> str:
        """检测段落样式名，返回对应标题 HTML 标签。

        支持中英文样式名：
            - Title / 标题 → <h1>
            - Heading 1 / 标题 1 → <h1>
            - Heading 2 / 标题 2 → <h2>
            - Heading 3 / 标题 3 → <h3>
            - Heading 4 / 标题 4 → <h4>

        Args:
            element: <w:p> XML 元素。
            text: 段落文本（已转义前的原始文本）。

        Returns:
            标题 HTML 字符串（如 "<h1>章节标题</h1>"），
            非标题样式返回空字符串。
        """
        try:
            # 从 <w:pPr>/<w:pStyle w:val="..."/> 提取样式 ID
            ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
            pPr = element.find(f"{ns}pPr")
            if pPr is None:
                return ""
            pStyle = pPr.find(f"{ns}pStyle")
            if pStyle is None:
                return ""

            style_id = pStyle.get(f"{ns}val", "")
            if not style_id:
                return ""

            # 样式 ID 通常为 Heading1/Heading2/Title 等
            # 也可能是中文名称的 hash，但常见情况下直接是英文名
            style_lower = style_id.lower()

            # 标题层级映射表
            heading_map = {
                "title": 1,
                "heading1": 1,
                "heading2": 2,
                "heading3": 3,
                "heading4": 4,
                # 中文样式名（部分文档使用）
                "1": 1,  # "标题 1" 的某些变体
            }

            level = heading_map.get(style_lower, 0)

            # 尝试匹配 HeadingN 模式（Heading5/Heading6 也支持）
            if level == 0 and style_lower.startswith("heading"):
                try:
                    level = int(style_lower[7:])
                    if level > 6:
                        level = 6
                except ValueError:
                    pass

            if level == 0:
                return ""

            escaped = DOCXParser._escape_html(text)
            return f"<h{level}>{escaped}</h{level}>"

        except Exception:
            return ""

    @staticmethod
    def _is_list_paragraph(element: Any) -> bool:
        """检测段落是否为列表项。

        列表段落包含 <w:numPr>（编号属性）或样式为 List Paragraph。

        Args:
            element: <w:p> XML 元素。

        Returns:
            True 如果段落是列表项。
        """
        try:
            ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
            pPr = element.find(f"{ns}pPr")
            if pPr is None:
                return False
            # 检测 <w:numPr>
            numPr = pPr.find(f"{ns}numPr")
            if numPr is not None:
                return True
            # 检测样式为 ListParagraph
            pStyle = pPr.find(f"{ns}pStyle")
            if pStyle is not None:
                style_id = pStyle.get(f"{ns}val", "").lower()
                if "list" in style_id:
                    return True
            return False
        except Exception:
            return False

    @staticmethod
    def _escape_html(text: str) -> str:
        """转义 HTML 特殊字符。"""
        if not text:
            return ""
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

    @staticmethod
    def _extract_table_html(element: Any) -> str:
        """从 <w:tbl> XML 元素提取表格并转为 HTML。

        使用 python-docx 的 Table 包装器获取行列数据。
        """
        try:
            from docx.table import Table

            table = Table(element, None)
            rows: list[list[str | None]] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            return DOCXParser._rows_to_html(rows)
        except Exception as exc:
            log.debug("docx.table_extract_failed", error=str(exc))
            return ""

    @staticmethod
    def _extract_headers_footers(docx_doc: Any) -> list[ParsedSection]:
        """提取 DOCX 页眉页脚文本。

        页眉页脚存储在独立的 section part 中，包含文档标题、
        页码、公司信息等元数据，对 RAG 检索有补充价值。

        Args:
            docx_doc: python-docx Document 对象。

        Returns:
            页眉页脚的 ParsedSection 列表。无页眉页脚或提取失败时返回空列表。
        """
        sections: list[ParsedSection] = []

        for section in docx_doc.sections:
            # 提取页眉
            try:
                header = section.header
                if header and not header.is_linked_to_previous:
                    header_text = "\n".join(
                        p.text.strip() for p in header.paragraphs if p.text.strip()
                    )
                    if header_text:
                        sections.append(
                            ParsedSection(
                                kind="text",
                                content=f"[页眉]\n{header_text}",
                                page=9998,
                            )
                        )
            except Exception as exc:
                log.debug("docx.header_extract_failed", error=str(exc))

            # 提取页脚
            try:
                footer = section.footer
                if footer and not footer.is_linked_to_previous:
                    footer_text = "\n".join(
                        p.text.strip() for p in footer.paragraphs if p.text.strip()
                    )
                    if footer_text:
                        sections.append(
                            ParsedSection(
                                kind="text",
                                content=f"[页脚]\n{footer_text}",
                                page=9999,
                            )
                        )
            except Exception as exc:
                log.debug("docx.footer_extract_failed", error=str(exc))

        return sections

    @staticmethod
    def _rows_to_html(rows: list[list[str | None]]) -> str:
        """将二维数据转为 HTML <table> 标签。

        第一行视为表头（<th>），其余为数据行（<td>）。
        """
        if not rows:
            return ""

        lines: list[str] = ["<table>"]

        for i, row in enumerate(rows):
            lines.append("<tr>")
            tag = "th" if i == 0 else "td"
            for cell in row:
                cell_text = (cell or "").strip()
                cell_text = (
                    cell_text.replace("&", "&amp;")
                    .replace("<", "&lt;")
                    .replace(">", "&gt;")
                )
                lines.append(f"<{tag}>{cell_text}</{tag}>")
            lines.append("</tr>")

        lines.append("</table>")
        return "\n".join(lines)

    async def _extract_images(
        self,
        docx_doc: Any,
        remaining: int,
        image_upload_enabled: bool = False,
        image_min_size: int = 50,
    ) -> tuple[list[ParsedSection], int]:
        """提取 DOCX 内嵌图片 — 支持上传 MinIO + VLM 描述 + 小图过滤。

        处理模式：
            - image_upload_enabled=True: 上传 MinIO → kind="image_url"
              + VLM 描述（VLM 可用时同时生成）
            - image_upload_enabled=False: 仅 VLM 描述 → kind="image_desc"
            - image_min_size: 宽或高小于此值的图片跳过

        Args:
            docx_doc: python-docx Document 对象。
            remaining: 本文档剩余可提取图片数。
            image_upload_enabled: 是否上传图片到 MinIO。
            image_min_size: 最小尺寸阈值（宽或高小于此值跳过）。

        Returns:
            (图片 ParsedSection 列表, 实际提取的图片数)
        """
        try:
            # 收集所有图片 part
            image_parts: list[tuple[bytes, str]] = []
            seen_rids: set[str] = set()

            for rid, rel in docx_doc.part.rels.items():
                if "image" not in rel.reltype.lower():
                    continue
                if rid in seen_rids:
                    continue
                seen_rids.add(rid)

                try:
                    blob = rel.target_part.blob
                    content_type = rel.target_part.content_type or "image/png"
                    if blob and len(blob) > 0:
                        image_parts.append((blob, content_type))
                except Exception as exc:
                    log.debug("docx.image_part_failed", rid=rid, error=str(exc))

            if not image_parts:
                return [], 0

            image_parts = image_parts[:remaining]

            semaphore = asyncio.Semaphore(_VLM_SEMAPHORE_LIMIT)

            async def process_image(
                idx: int, img_bytes: bytes, content_type: str
            ) -> ParsedSection | None:
                """处理单张图片 — 过滤、上传/描述。"""
                # 从 content_type 提取扩展名
                ext = content_type.split("/")[-1] if "/" in content_type else "png"
                if ext == "jpeg":
                    ext = "jpeg"
                elif ext == "jpg":
                    ext = "jpeg"

                # --- 小图过滤（仅在尺寸已知时过滤） ---
                if image_min_size > 0:
                    from app.document.image_storage import get_image_dimensions

                    img_width, img_height = get_image_dimensions(img_bytes, ext)
                    # 仅当尺寸确认大于 0 时才比较（未知尺寸的图片不过滤）
                    if img_width > 0 and img_height > 0:
                        if img_width < image_min_size or img_height < image_min_size:
                            log.debug(
                                "docx.images.filtered_small",
                                width=img_width,
                                height=img_height,
                                min_size=image_min_size,
                            )
                            return None

                # --- VLM 描述 ---
                vlm_desc = ""
                async with semaphore:
                    vlm_desc = await self._vlm_describe(img_bytes, content_type)

                desc_text = ""
                if vlm_desc and vlm_desc.strip():
                    desc_text = f"[图片描述: {vlm_desc.strip()}]"

                # --- 图片上传模式 ---
                if image_upload_enabled:
                    from app.document.image_storage import upload_image

                    url = await upload_image(
                        image_bytes=img_bytes,
                        ext=ext,
                        doc_id="docx",
                        page=0,
                        idx=idx,
                        min_size=0,  # 已过滤
                    )
                    if url:
                        return ParsedSection(
                            kind="image_url",
                            content=desc_text,
                            page=9999,  # 图片排在文本之后
                            image_url=url,
                        )
                    # 上传失败，降级为描述模式
                    if desc_text:
                        return ParsedSection(
                            kind="image_desc",
                            content=desc_text,
                            page=9999,
                        )
                    return None

                # --- 仅描述模式 ---
                if not desc_text:
                    return None
                return ParsedSection(
                    kind="image_desc",
                    content=desc_text,
                    page=9999,  # 图片排在文本之后
                )

            tasks = [
                process_image(idx, blob, ct)
                for idx, (blob, ct) in enumerate(image_parts)
            ]
            results = await asyncio.gather(*tasks, return_exceptions=True)

            sections: list[ParsedSection] = []
            count = 0
            for result in results:
                if isinstance(result, ParsedSection):
                    sections.append(result)
                    count += 1
                elif isinstance(result, Exception):
                    log.warning("docx.image_gather_error", error=str(result))

            return sections, count

        except Exception as exc:
            log.warning("docx.image_extract_error", error=str(exc))
            return [], 0

    async def _vlm_describe(self, image: bytes, mime_type: str) -> str:
        """调用 VLM 理解图片内容 — 延迟导入，不可用时返回空字符串。"""
        try:
            from app.vlm.provider import get_vision_provider

            vlm = get_vision_provider()
            return await vlm.understand(
                image=image,
                prompt=_IMAGE_PROMPT,
                mime_type=mime_type,
            )
        except ImportError:
            log.warning("docx.vlm_not_available")
            return ""
        except Exception as exc:
            log.warning("docx.vlm_error", error=str(exc))
            return ""
